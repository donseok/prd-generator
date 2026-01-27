"""
문서 파일(Word, PDF) 파서입니다.
외부 라이브러리(PyPDF2, python-docx)를 사용하여 텍스트를 추출합니다.
"""

from pathlib import Path
from typing import Optional

from app.models import InputType, ParsedContent, InputMetadata
from ..base_parser import BaseParser
from ..prompts.parsing_prompts import DOCUMENT_PARSING_PROMPT


class DocumentParser(BaseParser):
    """Word(.docx) 및 PDF(.pdf) 파일을 처리하는 파서입니다."""

    @property
    def supported_types(self) -> list[InputType]:
        return [InputType.DOCUMENT]

    @property
    def supported_extensions(self) -> list[str]:
        return [".pdf", ".docx", ".doc"]

    async def parse(
        self,
        file_path: Path,
        metadata: Optional[dict] = None
    ) -> ParsedContent:
        """파일 확장자에 따라 적절한 추출 함수를 호출합니다."""
        ext = file_path.suffix.lower()

        if ext == ".pdf":
            raw_text, pages = self._parse_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            raw_text, pages = self._parse_word(file_path)
        else:
            raise ValueError(f"지원하지 않는 문서 형식입니다: {ext}")

        # 메타데이터 생성
        file_metadata = await self.extract_metadata(file_path)
        file_metadata.page_count = len(pages)

        # 페이지별 섹션 생성
        sections = [
            {
                "title": f"페이지 {i+1}",
                "content": page,
            }
            for i, page in enumerate(pages)
        ]

        # 구조 데이터 생성
        structured_data = {
            "page_count": len(pages),
            "char_count": len(raw_text),
            "document_type": ext[1:].upper(),
        }

        # AI(Claude) 분석 (가능한 경우)
        if self.claude_client:
            try:
                analysis = await self._analyze_with_claude(raw_text)
                structured_data["ai_analysis"] = analysis
            except Exception as e:
                print(f"Claude 문서 분석 실패: {e}")

        return ParsedContent(
            raw_text=raw_text,
            structured_data=structured_data,
            metadata=file_metadata,
            sections=sections,
        )

    def _parse_pdf(self, file_path: Path) -> tuple[str, list]:
        """PyPDF2를 사용하여 PDF 내용을 텍스트로 추출합니다."""
        from PyPDF2 import PdfReader

        reader = PdfReader(str(file_path))
        pages = []
        all_text = []

        for page in reader.pages:
            text = page.extract_text() or ""
            pages.append(text)
            all_text.append(text)

        return "\n\n".join(all_text), pages

    def _parse_word(self, file_path: Path) -> tuple[str, list]:
        """python-docx를 사용하여 Word 문서를 추출합니다."""
        from docx import Document

        doc = Document(str(file_path))

        # 문단 추출
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 표 내용 추출
        tables_text = []
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(cells))
            tables_text.append("\n".join(table_rows))

        # 텍스트 합치기
        all_text = "\n\n".join(paragraphs)
        if tables_text:
            all_text += "\n\n=== 표 ===\n" + "\n\n".join(tables_text)

        # Word 문서는 전체를 하나의 페이지로 취급 (필요 시 섹션별 분리 가능)
        return all_text, [all_text]

    async def _analyze_with_claude(self, raw_text: str) -> dict:
        """Claude에게 문서 내용 요약 및 분석을 요청합니다."""
        result = await self.claude_client.complete_json(
            system_prompt=DOCUMENT_PARSING_PROMPT,
            user_prompt=f"다음 문서를 분석해주세요:\n\n{raw_text[:8000]}",
            temperature=0.2,
        )
        return result
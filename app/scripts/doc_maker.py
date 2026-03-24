#!/usr/bin/env python3
"""DOC (Word) maker script - converts MD documents to DOCX format.

Usage:
    python -m app.scripts.doc_maker                    # 전체 변환 (PRD+TRD+WBS+Proposal)
    python -m app.scripts.doc_maker --type prd         # PRD만 변환
    python -m app.scripts.doc_maker --type trd         # TRD만 변환
    python -m app.scripts.doc_maker --type wbs         # WBS만 변환
    python -m app.scripts.doc_maker --type proposal    # 제안서만 변환
"""

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

# Emoji-only removal pattern (preserves Korean punctuation and other valid characters)
_EMOJI_PATTERN = re.compile("["
    u"\U0001F600-\U0001F64F"  # emoticons
    u"\U0001F300-\U0001F5FF"  # symbols & pictographs
    u"\U0001F680-\U0001F6FF"  # transport & map
    u"\U0001F1E0-\U0001F1FF"  # flags
    u"\U00002702-\U000027B0"
    u"\U000024C2-\U0001F251"
    u"\U0001f926-\U0001f937"
    u"\U00010000-\U0010ffff"
    u"\u2640-\u2642"
    u"\u2600-\u2B55"
    u"\u200d"
    u"\u23cf"
    u"\u23e9"
    u"\u231a"
    u"\ufe0f"
    u"\u3030"
    "]+", flags=re.UNICODE)

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# 스타일 설정
FONT_NAME = '맑은 고딕'
COLORS = {
    'primary': RGBColor(0x7C, 0x3A, 0xED),    # 보라색
    'secondary': RGBColor(0x06, 0xB6, 0xD4),   # 시안
    'accent': RGBColor(0xF5, 0x9E, 0x0B),      # 주황
    'text': RGBColor(0x1A, 0x1A, 0x2E),         # 어두운 남색
    'text_light': RGBColor(0x64, 0x74, 0x8B),   # 회색
    'success': RGBColor(0x10, 0xB9, 0x81),      # 초록
    'warning': RGBColor(0xEF, 0x44, 0x44),      # 빨강
    'table_header_bg': RGBColor(0x7C, 0x3A, 0xED),
    'table_alt_bg': RGBColor(0xF5, 0xF3, 0xFF),
}


def setup_styles(doc: Document):
    """문서 기본 스타일 설정."""
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(11)
    font.color.rgb = COLORS['text']
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # 제목 스타일
    for level in range(1, 5):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_style.font.name = FONT_NAME
            h_style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            h_style.font.color.rgb = COLORS['primary']
            h_style.font.bold = True
            sizes = {1: 24, 2: 18, 3: 14, 4: 12}
            h_style.font.size = Pt(sizes.get(level, 12))
            h_style.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
            h_style.paragraph_format.space_after = Pt(6)


def add_cover_page(doc: Document, title: str, doc_type: str, metadata: dict = None):
    """표지 페이지 생성."""
    # 빈 줄 추가로 수직 중앙 배치 효과
    for _ in range(6):
        doc.add_paragraph()

    # 문서 타입 라벨
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_type)
    run.font.size = Pt(14)
    run.font.color.rgb = COLORS['secondary']
    run.font.name = FONT_NAME

    # 제목
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLORS['primary']
    run.font.name = FONT_NAME

    # 구분선
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 30)
    run.font.color.rgb = COLORS['text_light']
    run.font.size = Pt(12)

    # 메타데이터
    if metadata:
        for key, value in metadata.items():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'{key}: ')
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS['text_light']
            run.font.name = FONT_NAME
            run = p.add_run(str(value))
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS['text']
            run.font.name = FONT_NAME
            run.font.bold = True

    # 생성일
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n생성일: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    run.font.size = Pt(10)
    run.font.color.rgb = COLORS['text_light']
    run.font.name = FONT_NAME

    doc.add_page_break()


def set_cell_shading(cell, color_hex: str):
    """셀 배경색 설정."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc: Document, headers: list, rows: list):
    """스타일이 적용된 테이블 추가."""
    if not headers or not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 헤더
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = FONT_NAME
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, '7C3AED')

    # 데이터 행
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            run.font.size = Pt(10)
            run.font.name = FONT_NAME
            run.font.color.rgb = COLORS['text']
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'F5F3FF')

    doc.add_paragraph()  # 테이블 뒤 간격


def parse_markdown_table(lines: list) -> tuple:
    """마크다운 테이블을 파싱하여 (headers, rows) 반환."""
    if len(lines) < 2:
        return None, None

    # 헤더 파싱
    header_line = lines[0].strip()
    if not header_line.startswith('|'):
        return None, None

    headers = [h.strip() for h in header_line.strip('|').split('|')]
    headers = [re.sub(r'\*\*(.+?)\*\*', r'\1', h) for h in headers]

    # 구분선 건너뛰기
    data_start = 1
    if len(lines) > 1 and re.match(r'\s*\|[\s\-:|]+\|', lines[1]):
        data_start = 2

    # 데이터 행
    rows = []
    for line in lines[data_start:]:
        line = line.strip()
        if not line.startswith('|'):
            break
        cells = [c.strip() for c in line.strip('|').split('|')]
        cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c) for c in cells]
        rows.append(cells)

    return headers, rows


def md_to_docx(md_content: str, doc_type: str, title: str = None) -> Document:
    """마크다운 콘텐츠를 DOCX 문서로 변환."""
    doc = Document()

    # 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    setup_styles(doc)

    # 제목 추출
    if not title:
        first_heading = re.search(r'^#\s+(.+)$', md_content, re.MULTILINE)
        title = first_heading.group(1) if first_heading else f'{doc_type} 문서'
    title = re.sub(r'\*\*(.+?)\*\*', r'\1', title).strip()

    # 표지
    type_labels = {
        'PRD': 'Product Requirements Document',
        'TRD': 'Technical Requirements Document',
        'WBS': 'Work Breakdown Structure',
        'PROP': '고객 제안서 (Proposal)',
    }
    add_cover_page(doc, title, type_labels.get(doc_type, doc_type))

    # 마크다운 파싱 및 변환
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # 코드 블록 처리
        if line.strip().startswith('```'):
            if in_code_block:
                # 코드 블록 종료
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = COLORS['text']
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 테이블 처리
        if line.strip().startswith('|'):
            table_lines.append(line)
            i += 1
            continue
        elif table_lines:
            # 테이블 종료
            headers, rows = parse_markdown_table(table_lines)
            if headers and rows:
                add_styled_table(doc, headers, rows)
            table_lines = []

        stripped = line.strip()

        # 빈 줄
        if not stripped:
            i += 1
            continue

        # 수평선
        if re.match(r'^-{3,}$|^\*{3,}$|^_{3,}$', stripped):
            i += 1
            continue

        # 제목
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            level = min(level, 4)
            doc.add_heading(text, level=level)
            i += 1
            continue

        # 블록 인용
        if stripped.startswith('>'):
            quote_text = re.sub(r'^>\s*', '', stripped)
            quote_text = re.sub(r'\*\*(.+?)\*\*', r'\1', quote_text)
            # 이모지 제거
            quote_text = _EMOJI_PATTERN.sub('', quote_text).strip()
            if quote_text:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.0)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(quote_text)
                run.font.italic = True
                run.font.color.rgb = COLORS['text_light']
                run.font.name = FONT_NAME
                run.font.size = Pt(10)
            i += 1
            continue

        # 불릿 리스트
        list_match = re.match(r'^(\s*)([-*+]|\d+[.)]) (.+)$', stripped)
        if list_match:
            text = list_match.group(3)
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
            # 이모지 정리
            text = _EMOJI_PATTERN.sub('', text).strip()
            if text:
                p = doc.add_paragraph(style='List Bullet')
                # Bold 처리
                bold_parts = re.split(r'(`[^`]+`)', text)
                for part in bold_parts:
                    if part.startswith('`') and part.endswith('`'):
                        run = p.add_run(part[1:-1])
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                    else:
                        run = p.add_run(part)
                        run.font.name = FONT_NAME
                        run.font.size = Pt(10)
                run.font.color.rgb = COLORS['text']
            i += 1
            continue

        # 일반 텍스트
        text = stripped
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = _EMOJI_PATTERN.sub('', text).strip()
        if text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run.font.size = Pt(11)
            run.font.color.rgb = COLORS['text']
        i += 1

    # 남은 테이블 처리
    if table_lines:
        headers, rows = parse_markdown_table(table_lines)
        if headers and rows:
            add_styled_table(doc, headers, rows)

    return doc


def find_latest_md(output_dir: Path, prefix: str) -> Path | None:
    """최신 MD 파일 찾기."""
    md_files = list(output_dir.glob(f'{prefix}-*.md'))
    if not md_files:
        return None
    return max(md_files, key=lambda x: x.stat().st_mtime)


def convert_single(doc_type: str) -> Path | None:
    """단일 문서 타입 변환."""
    type_config = {
        'prd': {'dir': 'workspace/outputs/prd', 'prefix': 'PRD', 'label': 'PRD'},
        'trd': {'dir': 'workspace/outputs/trd', 'prefix': 'TRD', 'label': 'TRD'},
        'wbs': {'dir': 'workspace/outputs/wbs', 'prefix': 'WBS', 'label': 'WBS'},
        'proposal': {'dir': 'workspace/outputs/proposals', 'prefix': 'PROP', 'label': 'PROP'},
    }

    config = type_config.get(doc_type)
    if not config:
        print(f'알 수 없는 문서 타입: {doc_type}')
        return None

    output_dir = Path(config['dir'])
    md_path = find_latest_md(output_dir, config['prefix'])

    if not md_path:
        print(f'{config["label"]} MD 파일을 찾을 수 없습니다. ({output_dir})')
        return None

    print(f'  입력: {md_path.name}')
    md_content = md_path.read_text(encoding='utf-8')

    doc = md_to_docx(md_content, config['label'])

    # DOCX 저장 (같은 디렉토리, 같은 타임스탬프)
    stem = md_path.stem  # e.g., PRD-20260305-120000
    docx_dir = Path('workspace/outputs/doc')
    docx_dir.mkdir(parents=True, exist_ok=True)
    docx_path = docx_dir / f'{stem}.docx'
    doc.save(str(docx_path))
    print(f'  출력: {docx_path}')

    return docx_path


def main():
    parser = argparse.ArgumentParser(description="문서를 Word(DOCX) 형식으로 변환")
    parser.add_argument(
        "--type", "-t",
        type=str,
        choices=['prd', 'trd', 'wbs', 'proposal', 'all'],
        default='all',
        help="변환할 문서 타입 (기본: all)"
    )
    args = parser.parse_args()

    print('\n' + '=' * 70)
    print('Word(DOCX) 문서 변환')
    print(f'시작 시간: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    print('=' * 70)

    doc_types = ['prd', 'trd', 'wbs', 'proposal'] if args.type == 'all' else [args.type]
    results = {}

    for doc_type in doc_types:
        print(f'\n[{doc_type.upper()}] 변환 중...')
        result = convert_single(doc_type)
        results[doc_type] = result

    # 결과 요약
    print('\n' + '=' * 70)
    print('변환 결과')
    print('=' * 70)

    success_count = 0
    for doc_type, path in results.items():
        status = 'OK' if path else 'SKIP (MD 파일 없음)'
        if path:
            success_count += 1
        print(f'  {doc_type.upper():10s}: {status}')

    print(f'\n  성공: {success_count}/{len(results)}')
    print(f'  출력 디렉토리: workspace/outputs/doc/')


if __name__ == "__main__":
    main()
